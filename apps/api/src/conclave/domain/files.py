from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from conclave.config import settings
from conclave.db.session import DATA_DIR
from conclave.domain.redact import redact_pii

ALLOWED_SUFFIXES = {".md", ".txt", ".csv", ".json", ".pdf", ".docx"}

_ocr_engine = None


def _get_ocr():
    """Lazy singleton: RapidOCR loads its ONNX models once per process."""
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR

        _ocr_engine = RapidOCR()
    return _ocr_engine


def conversation_dir(conversation_id: str) -> Path:
    path = DATA_DIR / "conversations" / conversation_id
    path.mkdir(parents=True, exist_ok=True)
    (path / "files").mkdir(exist_ok=True)
    return path


def remove_conversation_dir(conversation_id: str) -> None:
    path = DATA_DIR / "conversations" / conversation_id
    if path.exists():
        shutil.rmtree(path, ignore_errors=True)


def shared_doc_path(conversation_id: str) -> Path:
    return conversation_dir(conversation_id) / "shared.md"


def read_shared_doc(conversation_id: str) -> str:
    path = shared_doc_path(conversation_id)
    if not path.exists():
        path.write_text("# Shared document\n\n", encoding="utf-8")
    return path.read_text(encoding="utf-8")


def write_shared_doc(conversation_id: str, content: str) -> str:
    path = shared_doc_path(conversation_id)
    path.write_text(content, encoding="utf-8")
    return content


def _ocr_pdf_text(p: Path) -> str:
    """OCR an image-only PDF locally (RapidOCR — nothing leaves the machine)."""
    try:
        import numpy as np
        import pypdfium2 as pdfium

        ocr = _get_ocr()
        pdf = pdfium.PdfDocument(str(p))
        try:
            limit = settings.ocr_max_pages or len(pdf)
            pages: list[str] = []
            for i in range(len(pdf)):
                if i >= limit:
                    pages.append(f"…[{len(pdf) - limit} more pages not OCRed]")
                    break
                bitmap = pdf[i].render(scale=2.0)  # ~144 dpi: decent recognition, sane speed
                result, _ = ocr(np.array(bitmap.to_pil()))
                if result:
                    # Page markers keep chronology legible across a long record.
                    pages.append(f"[page {i + 1}]\n" + "\n".join(item[1] for item in result))
            return "\n\n".join(pages).strip()
        finally:
            pdf.close()
    except Exception:  # noqa: BLE001 — OCR failure degrades to the no-text note
        return ""


def _extract_pdf_text(p: Path) -> tuple[str, str]:
    """Returns (text, method) where method is pdf-text | pdf-ocr | empty | failed."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        pages: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        text = "\n\n".join(pages).strip()
        if text:
            return text, "pdf-text"
        ocr_text = _ocr_pdf_text(p)
        if ocr_text:
            return "[OCR — may contain recognition errors]\n" + ocr_text, "pdf-ocr"
        return "[PDF contains no extractable text — likely a scanned image]", "empty"
    except Exception:  # noqa: BLE001 — malformed uploads must not break a turn
        return "[PDF could not be parsed]", "failed"


def _extract_docx_text(p: Path) -> tuple[str, str]:
    """Returns (text, method) where method is docx | empty | failed."""
    try:
        import docx

        d = docx.Document(str(p))
        parts = [para.text for para in d.paragraphs if para.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts).strip()
        if text:
            return text, "docx"
        return "[Word document contains no extractable text]", "empty"
    except Exception:  # noqa: BLE001 — malformed uploads must not break a turn
        return "[Word document could not be parsed]", "failed"


@dataclass
class Extraction:
    """What a turn would actually see for an attachment, plus how it was obtained."""

    text: str
    method: str  # text | pdf-text | pdf-ocr | docx | empty | failed | missing
    chars: int

    @property
    def usable(self) -> bool:
        return self.method not in ("empty", "failed", "missing") and self.chars > 0


def extract_attachment(path: str) -> Extraction:
    p = Path(path)
    if not p.exists():
        return Extraction("", "missing", 0)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        text, method = _extract_pdf_text(p)
    elif suffix == ".docx":
        text, method = _extract_docx_text(p)
    else:
        try:
            text, method = p.read_text(encoding="utf-8", errors="replace"), "text"
        except OSError:
            return Extraction("", "failed", 0)
        if not text.strip():
            method = "empty"
    if settings.redact_pii:
        text = redact_pii(text)
    cap = settings.attachment_max_chars
    if cap and len(text) > cap:
        text = text[:cap] + f"\n…[truncated at {cap} chars — raise CONCLAVE_ATTACHMENT_MAX_CHARS]"
    return Extraction(text, method, len(text))


def read_attachment_text(path: str) -> str:
    """Turn-facing read: the text an expert sees."""
    return extract_attachment(path).text
