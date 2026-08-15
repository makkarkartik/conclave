"""Attachment extraction: PDF and DOCX must yield real text for the
read_attachment tool; text formats pass through with the size cap."""

from __future__ import annotations

from pathlib import Path

from conclave.domain.files import MAX_READ_CHARS, read_attachment_text


def _minimal_pdf(text: str) -> bytes:
    """A hand-built single-page PDF with one Helvetica text run."""
    stream = f"BT /F1 12 Tf 50 700 Td ({text}) Tj ET".encode()
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objs) + 1}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        b"trailer\n<< /Size " + str(len(objs) + 1).encode() + b" /Root 1 0 R >>\n"
        b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return bytes(out)


def test_pdf_text_is_extracted(tmp_path: Path):
    pdf = tmp_path / "brief.pdf"
    pdf.write_bytes(_minimal_pdf("Deliberation payload 42"))
    text = read_attachment_text(str(pdf))
    assert "Deliberation payload 42" in text


def test_broken_pdf_degrades_gracefully(tmp_path: Path):
    pdf = tmp_path / "junk.pdf"
    pdf.write_bytes(b"%PDF-1.4 this is not really a pdf")
    text = read_attachment_text(str(pdf))
    assert "could not be parsed" in text or "no extractable text" in text


def test_docx_paragraphs_and_tables_are_extracted(tmp_path: Path):
    import docx

    d = docx.Document()
    d.add_paragraph("Quarterly risk memo")
    d.add_paragraph("The canary rollout must be gated.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "metric"
    table.rows[0].cells[1].text = "threshold"
    path = tmp_path / "memo.docx"
    d.save(str(path))

    text = read_attachment_text(str(path))
    assert "Quarterly risk memo" in text
    assert "canary rollout" in text
    assert "metric | threshold" in text


def test_text_passthrough_and_cap(tmp_path: Path):
    f = tmp_path / "notes.md"
    f.write_text("x" * (MAX_READ_CHARS + 500), encoding="utf-8")
    text = read_attachment_text(str(f))
    assert text.endswith("…[truncated]")
    assert len(text) <= MAX_READ_CHARS + 20
